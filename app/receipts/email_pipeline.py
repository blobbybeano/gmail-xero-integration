"""Email Invoice Importer — scan pipeline.

Stateless helpers plus ``scan_email_batch()`` (the main entry point).
No Flask dependency — safe to call from the background scheduler in main.py
or from a web-request background thread.
"""

from __future__ import annotations

import hashlib
import json
import logging
import re
import threading
from typing import Any

from .email_store import (
    STATUS_DUPLICATE,
    STATUS_IGNORED,
    STATUS_IMPORTED,
    STATUS_NEW,
    STATUS_NO_ACCOUNT,
    STATUS_NOT_INVOICE,
    STATUS_OWN_COMPANY,
    STATUS_POSSIBLE_DUP,
    STATUS_SKIPPED_EMAIL,
    STATUS_SUSPICIOUS,
    create_item,
    get_batch,
    list_items,
    message_already_scanned,
    update_batch,
    update_item,
)
from ..admin_store import get_json_setting
from .expense_store import list_all_receipts, create_receipt

_log = logging.getLogger(__name__)
_ACCOUNT_LEARNING_KEY = "account_category_learning"

# One scan at a time (shared across web + scheduler calls)
_scan_lock = threading.Lock()


# ── Merchant normalisation ────────────────────────────────────────────────────

def normalize_merchant(name: str) -> str:
    """Lowercase, alphanumeric-only for fuzzy merchant matching."""
    return re.sub(r"[^a-z0-9]", "", (name or "").lower())


def derive_supplier_merchant(
    merchant: str,
    *,
    from_name: str,
    from_addr: str,
    subject: str,
    own_names: list[str] | None,
) -> str:
    """Fix the classic supplier-invoice OCR mistake: invoices addressed TO us
    often get OUR OWN company name extracted as the 'merchant' (it is usually
    the most prominent name on the page). That makes the item unrecognisable
    in the review list ("POW Services Limited £126" instead of "Redwood Wales
    / CJH £126") and can mislead the AI account coding.

    If the OCR'd merchant looks like one of our own names, derive the real
    supplier instead, in order of preference:
      1. subject pattern "... from <Supplier> [for <us>]"
      2. the sender's display name
      3. the sender's email domain (e.g. racbusinessclub.co.uk)
    Falls back to the original merchant if nothing better is found.
    """
    own_norms = [normalize_merchant(n) for n in (own_names or [])]
    own_norms = [n for n in own_norms if n]

    def _is_own(text: str) -> bool:
        t = normalize_merchant(text)
        if not t:
            return False
        return any(o in t or t in o for o in own_norms)

    if not merchant or not _is_own(merchant):
        return merchant

    m = re.search(r"\bfrom\s+(.+?)(?=\s+for\s+|\s*$)", subject or "", re.IGNORECASE)
    if m:
        cand = m.group(1).strip(" .,-–—")
        if cand and not _is_own(cand):
            return cand

    if from_name and not _is_own(from_name):
        return from_name.strip()

    dom = (from_addr or "").rsplit("@", 1)[-1].strip().lower()
    if dom and "." in dom and not _is_own(dom):
        return dom

    return merchant


# ── Own-company filter ────────────────────────────────────────────────────────

# Xero sends invoice copies back to you from these domains when you raise and
# send an invoice to a client.  The subject typically reads "Invoice from
# Power Wash Ltd" — i.e. the own-company name appears in the *subject*, not
# the sender, because the actual sender is Xero's mail infrastructure.
_XERO_GATEWAY_DOMAINS = (
    "post.xero.com",
    "xero.com",
    "invoices.xero.com",
)


def is_own_company_sender(
    from_addr: str,
    from_name: str,
    subject: str = "",
    *,
    own_names: list[str],
    own_domains: list[str],
) -> bool:
    """Return True when this email is one of our *own outgoing* invoices.

    Three detection paths:

    1. **Own domain** — the From address contains our own email domain
       (e.g. powwash.co.uk).

    2. **Own display name** — the From display name / address contains our
       company name (e.g. "Power Wash Ltd" or "powwash" in the header name
       or address). Catches cases where a third-party tool sends on our behalf
       but preserves our name as the sender.

    3. **"from [our company]" in subject** — works for any sender. Our
       outgoing invoice emails say "Invoice from Power Wash Ltd" or
       "Invoice INV-0042 from Pow Services Limited for Rory Kilgour",
       including customer reply threads ("Re: Invoice INV-5648 from Pow
       Services Limited"). Supplier invoices *addressed to* us say "Invoice
       for/to Power Wash Ltd" — the preposition "for/to" prevents false
       matches. Also covers Xero invoice-copy gateway emails.
    """
    addr = (from_addr or "").lower()
    name = (from_name or "").lower()
    subj = (subject or "").lower()

    # Path 1: own domain in sender address
    for d in own_domains:
        dl = d.lower().strip()
        if dl and dl in addr:
            return True

    # Path 2: own company name in sender display name or address
    for n in own_names:
        nl = n.lower().strip()
        if nl and (nl in name or nl in addr):
            return True

    # Path 3: subject says "from [our company]" — regardless of who sent it.
    # Our outgoing invoices read: "Invoice from Power Wash Ltd" or
    # "Invoice INV-0042 from Power Wash Ltd for Rory Kilgour" — including
    # customer REPLY threads ("Re: Invoice INV-5648 from Pow Services Limited").
    # Supplier invoices addressed TO us say "Invoice for/to Power Wash Ltd"
    # not "from", so this check safely distinguishes the two cases.
    # Also covers Xero invoice-copy gateway emails (noreply@post.xero.com).
    for n in own_names:
        nl = n.lower().strip()
        if nl and f"from {nl}" in subj:
            return True

    return False


# ── Attachment filter ─────────────────────────────────────────────────────────

_PHOTO_RE = re.compile(
    r"^(img_|dsc_|dsc\d|dcim|pano_|vid_|mov_|photo|screenshot|screen_shot)",
    re.IGNORECASE,
)
_INVOICE_EXTS  = {".pdf", ".jpg", ".jpeg", ".png", ".gif", ".webp", ".tiff"}
_INVOICE_MIMES = {
    "application/pdf", "image/jpeg", "image/jpg", "image/png",
    "image/gif", "image/webp", "image/tiff",
}


def is_invoice_attachment(filename: str, mime_type: str) -> bool:
    """Heuristic: PDF or invoice-shaped image, not a camera shot."""
    fn  = (filename or "").strip()
    mt  = (mime_type or "").lower().split(";")[0].strip()
    ext = ("." + fn.rsplit(".", 1)[-1]).lower() if "." in fn else ""
    if ext not in _INVOICE_EXTS and mt not in _INVOICE_MIMES:
        return False
    if _PHOTO_RE.match(fn):
        return False
    return True


# ── Amount reconciliation ─────────────────────────────────────────────────────

def reconcile_amounts(
    total: float | None,
    net: float | None,
    tax: float | None,
    *,
    vat_rate: float = 20.0,
) -> tuple[float | None, float | None, float | None]:
    """Return (inc, ex, vat) with best-effort fill-in."""
    if total is None and net is None:
        return None, None, None
    if total is not None:
        inc = round(float(total), 2)
        supplied_net = round(float(net), 2) if net is not None else None
        supplied_tax = round(float(tax), 2) if tax is not None else None

        # OCR sometimes lifts a VAT/tax figure from the wrong part of a PDF
        # (for example a statement summary) while the invoice total is correct.
        # Never let an impossible positive VAT amount make a positive invoice
        # have a negative net amount.
        if inc > 0:
            if supplied_tax is not None and (supplied_tax < 0 or supplied_tax > inc):
                supplied_tax = None
            if supplied_net is not None and (supplied_net < 0 or supplied_net > inc):
                supplied_net = None

        if supplied_net is not None:
            ex = supplied_net
            vat = round(inc - ex, 2)
        elif supplied_tax is not None:
            vat = supplied_tax
            ex = round(inc - vat, 2)
        else:
            ex  = round(inc / (1 + vat_rate / 100), 2)
            vat = round(inc - ex, 2)
    else:
        ex  = round(float(net), 2)
        vat = round(float(tax), 2) if tax is not None else round(ex * vat_rate / 100, 2)
        inc = round(ex + vat, 2)
    return inc, ex, vat


_MONEY_RE = r"([+-]?\d{1,6}(?:,\d{3})*(?:\.\d{2})?)"


def _to_money(value: str) -> float | None:
    try:
        return round(float((value or "").replace(",", "")), 2)
    except (TypeError, ValueError):
        return None


def explicit_total_from_text(raw_text: str) -> float | None:
    """Return an explicit document-level total from OCR text when present.

    Document AI sometimes extracts the first row total from a table instead of
    the final amount due.  Supplier PDFs often print the real total with a
    phrase like "Total including all applicable taxes"; prefer those labelled
    final totals over table line totals.
    """
    text = raw_text or ""
    strong_patterns = [
        r"total\s+including\s+all\s+applicable\s+taxes\s*(?:[:\-])?\s*(?:GBP|£)?\s*" + _MONEY_RE,
        r"total\s+(?:including|incl\.?|inc\.?)\s+(?:vat|tax(?:es)?)\s*(?:[:\-])?\s*(?:GBP|£)?\s*" + _MONEY_RE,
        r"(?:balance|amount)\s+due\s*(?:[:\-])?\s*(?:GBP|£)?\s*" + _MONEY_RE,
        r"grand\s+total\s*(?:[:\-])?\s*(?:GBP|£)?\s*" + _MONEY_RE,
        r"invoice\s+total\s*(?:[:\-])?\s*(?:GBP|£)?\s*" + _MONEY_RE,
        r"total\s+(?:payable|paid)\s*(?:[:\-])?\s*(?:GBP|£)?\s*" + _MONEY_RE,
    ]
    for pat in strong_patterns:
        matches = re.findall(pat, text, flags=re.IGNORECASE)
        vals = [_to_money(m) for m in matches]
        vals = [v for v in vals if v is not None]
        if vals:
            return vals[-1]

    # Weaker fallback: if the word "Total" appears several times, the last
    # total-like amount is normally the document total, not an earlier row.
    weak = re.findall(
        r"\btotal\b\s*(?:[:\-])?\s*(?:GBP|£)?\s*" + _MONEY_RE,
        text,
        flags=re.IGNORECASE,
    )
    vals = [_to_money(m) for m in weak]
    vals = [v for v in vals if v is not None]
    return vals[-1] if vals else None


def reconcile_email_amounts_from_text(
    total: float | None,
    net: float | None,
    tax: float | None,
    raw_text: str,
    *,
    vat_rate: float = 20.0,
) -> tuple[float | None, float | None, float | None]:
    """Reconcile email invoice amounts with a final-total OCR sanity pass."""
    explicit_total = explicit_total_from_text(raw_text)
    if explicit_total is not None:
        current_total = _to_money(str(total)) if total is not None else None
        # Prefer a labelled final total when OCR picked a smaller table row.
        if current_total is None or abs(explicit_total - current_total) > 0.01:
            total = explicit_total
            supplied_tax = _to_money(str(tax)) if tax is not None else None
            supplied_net = _to_money(str(net)) if net is not None else None
            # If the old "tax" field is actually the same as the explicit
            # final total, it was not a VAT amount. Drop it.
            if supplied_tax is not None and abs(supplied_tax - explicit_total) <= 0.01:
                tax = None
            if (
                supplied_net is not None
                and supplied_tax is not None
                and abs((supplied_net + supplied_tax) - explicit_total) > 0.02
            ):
                net = None
                tax = None

    lowered = (raw_text or "").lower()
    if explicit_total is not None and (
        "insurance premium tax" in lowered
        or re.search(r"\bipt\b", lowered)
        or "premiums include ipt" in lowered
    ):
        # IPT is not VAT.  For RAC/insurance notices, record the full gross as
        # no-VAT so the app does not invent reclaimable VAT.
        return round(float(explicit_total), 2), round(float(explicit_total), 2), 0.0

    return reconcile_amounts(total, net, tax, vat_rate=vat_rate)


# ── SHA-256 ───────────────────────────────────────────────────────────────────

def sha256_hex(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# ── Word-document (.docx) invoices ────────────────────────────────────────────

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def is_word_doc(filename: str, mime_type: str) -> bool:
    """True for Word .docx files (some suppliers send invoices as Word, not PDF)."""
    mt = (mime_type or "").lower().split(";")[0].strip()
    fn = (filename or "").lower()
    return mt == _DOCX_MIME or fn.endswith(".docx")


def extract_docx_text(data: bytes) -> str:
    """Extract readable text from a .docx — paragraphs AND table cells.

    Invoice totals are almost always inside tables, which python-docx does not
    include in ``paragraphs``, so we walk both. Returns "" if unreadable.
    """
    import io
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


# ── Smart categorisation from history ────────────────────────────────────────

def smart_categorise(
    merchant_norm: str,
    *,
    exp_receipts: list[dict],
    email_items: list[dict],
) -> tuple[str, str]:
    """Return (account_code, account_name) from prior submissions for this merchant.

    Checks in-batch email items first (most recently added last → reversed),
    then expense_receipts (newest first).
    """
    if not merchant_norm:
        return "", ""

    for it in reversed(email_items):
        if (
            normalize_merchant(it.get("merchant", "")) == merchant_norm
            and it.get("category_account_code")
            and it["status"] not in (STATUS_DUPLICATE, STATUS_IGNORED, STATUS_OWN_COMPANY)
        ):
            return it["category_account_code"], it.get("category_account_name", "")

    for r in exp_receipts:
        m = normalize_merchant(
            r.get("merchant") or r.get("ocr_merchant") or ""
        )
        if m == merchant_norm and r.get("category_account_code"):
            return r["category_account_code"], r.get("category_account_name", "")

    return "", ""


def _account_learning_key(merchant: str) -> str:
    key = re.sub(r"[^a-z0-9]+", "", (merchant or "").lower())
    for suffix in ("limited", "ltd", "uk", "gb"):
        if key.endswith(suffix) and len(key) > len(suffix) + 3:
            key = key[: -len(suffix)]
            break
    return key


def learned_categorise(db_path: str, merchant: str) -> tuple[str, str]:
    key = _account_learning_key(merchant)
    if not key:
        return "", ""
    try:
        learned = get_json_setting(db_path, _ACCOUNT_LEARNING_KEY, {}) or {}
    except Exception:
        learned = {}
    if not isinstance(learned, dict):
        return "", ""
    hit = learned.get(key) or {}
    if not isinstance(hit, dict):
        return "", ""
    return str(hit.get("code") or "").strip(), str(hit.get("name") or "").strip()


def _norm_account(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def _find_account_by_terms(
    accounts: list[dict],
    *,
    any_terms: tuple[str, ...],
    exclude_terms: tuple[str, ...] = (),
) -> tuple[str, str]:
    best: tuple[int, str, str] | None = None
    for a in accounts or []:
        code = str(a.get("Code") or "").strip()
        name = str(a.get("Name") or "").strip()
        if not code:
            continue
        norm = _norm_account(name)
        if exclude_terms and any(t in norm for t in exclude_terms):
            continue
        hits = sum(1 for t in any_terms if t in norm)
        if not hits:
            continue
        score = hits * 10
        if "vehicle" in norm or "motor" in norm or "van" in norm:
            score += 4
        if "advertis" in norm or "marketing" in norm:
            score += 4
        if "bank" in norm or "merchant" in norm:
            score += 4
        if best is None or score > best[0]:
            best = (score, code, name)
    return ("", "") if best is None else (best[1], best[2])


def _rule_text(merchant: str, raw_text: str) -> str:
    return " ".join((merchant or "", raw_text or "")).lower()


def rule_based_categorise(
    merchant: str,
    raw_text: str,
    exp_accounts: list[dict],
) -> tuple[str, str]:
    text = _rule_text(merchant, raw_text)
    if not text or not exp_accounts:
        return "", ""

    if any(t in text for t in (
        "checkatrade", "check a trade", "lead generation", "google ads",
        "google adwords", "facebook ads", "meta ads", "bark com", "rated people",
        "yell", "directory listing", "sponsored listing", "lead fee",
    )):
        return _find_account_by_terms(
            exp_accounts,
            any_terms=("advertis", "marketing", "lead", "promotion"),
            exclude_terms=("clean", "material", "fuel", "vehicle", "motor"),
        )

    if any(t in text for t in (
        "tender pos", "tenderpos", "card terminal", "card transaction",
        "merchant service", "merchant services", "payment processing",
        "terminal rental", "stripe", "sumup", "sum up", "zettle", "izettle",
        "squareup", "square ", "worldpay", "paypal fees", "cashflows fee",
        "cashflows fees", "card processing",
    )):
        code, name = _find_account_by_terms(
            exp_accounts,
            any_terms=(
                "merchant", "bank charge", "bank charges", "card fee",
                "card fees", "payment processing", "transaction fee",
                "fees", "charges",
            ),
            exclude_terms=("software", "computer", "it ", "material", "clean"),
        )
        if code:
            return code, name
        return _find_account_by_terms(
            exp_accounts,
            any_terms=("bank", "finance charge", "charges"),
            exclude_terms=("software", "computer", "it ", "material", "clean"),
        )

    if "rac" in text and any(t in text for t in ("breakdown", "business club", "vehicle", "cover")):
        return _find_account_by_terms(
            exp_accounts,
            any_terms=("vehicle", "motor", "van", "travel"),
            exclude_terms=("fuel", "parking", "insurance"),
        )

    if any(t in text for t in (
        "ringgo", "ring go", "parking", "car park", "ncp", "paybyphone",
        "justpark", "residents permit", "resident permit", "parking permit",
        "visitor permit", "cpz", "controlled parking zone", "merton council",
    )):
        return _find_account_by_terms(
            exp_accounts,
            any_terms=("vehicle", "motor", "van", "travel"),
            exclude_terms=("fuel", "rates", "rateable", "business rates", "council tax"),
        )

    if any(t in text for t in (
        "eca cleaning", "cleaning supplies", "window cleaning warehouse",
        "softwash", "hypo", "sodium hypochlorite", "biocide",
        "cleaning solution", "pressure washing chemical", "gutter joint",
        "gutter seal",
    )):
        return _find_account_by_terms(
            exp_accounts,
            any_terms=("material", "materials", "tools", "consumable", "supplies"),
            exclude_terms=("fuel", "vehicle", "motor", "parking", "clean"),
        )

    if any(t in text for t in (
        "microsoft", "office 365", "google workspace", "google cloud",
        "adobe", "openai", "github", "replit", "fly io", "fly.io",
        "digitalocean", "aws", "amazon web services", "dropbox",
        "software subscription", "saas",
    )):
        return _find_account_by_terms(
            exp_accounts,
            any_terms=("software", "computer", "it ", "subscription"),
            exclude_terms=("merchant", "bank charge", "vehicle", "fuel", "material", "clean"),
        )

    return "", ""


def tax_computation_not_supplier_invoice(merchant: str, raw_text: str) -> str:
    """Return a reason when this is a tax-return/computation, not a supplier bill.

    Accountants often send documents showing how much is owed to HMRC. Those can
    contain amounts, due dates and payment wording, but they are not a charge
    for the accountant's services and should not be imported as an expense.
    """
    text = _rule_text(merchant, raw_text)
    if not text:
        return ""
    tax_terms = (
        "tax computation", "corporation tax computation", "income tax computation",
        "self assessment tax return", "tax return", "ct600", "sa302",
        "hmrc", "corporation tax due", "tax payable", "amount due to hmrc",
        "payment due to hmrc", "tax calculation",
    )
    service_fee_terms = (
        "invoice number", "invoice no", "tax invoice", "vat invoice",
        "professional fee", "professional fees", "accountancy fee",
        "accountancy fees", "our fee", "our fees", "services provided",
        "fee note",
    )
    if any(t in text for t in tax_terms) and not any(t in text for t in service_fee_terms):
        return "Tax computation/return document, not supplier charge"
    return ""


# ── Deduplication ─────────────────────────────────────────────────────────────

def dedup_against_receipts(
    content_sha: str,
    merchant_norm: str,
    purchased_on: str,
    amount_inc: float | None,
    existing_receipts: list[dict],
    existing_scan_items: list[dict],
) -> tuple[str, str, str]:
    """Three-layer dedup: exact SHA → logical match in receipts → logical in batch.

    Returns (status, dup_reason, match_id).
    """
    sha_prefix = content_sha[:16]

    # 1 — exact image match in prior email scan items
    for it in existing_scan_items:
        sf = it.get("stored_file", "")
        if sf and sha_prefix in sf:
            return STATUS_DUPLICATE, "Identical attachment already scanned", it["id"]

    # 2 — exact image match in expense_receipts
    for r in existing_receipts:
        sf = r.get("stored_file", "")
        if sf and sha_prefix in sf:
            return STATUS_DUPLICATE, "Identical image already in Field Expenses", r["id"]

    # 3 — logical match (merchant + date + amount ±£0.02)
    if merchant_norm and purchased_on and amount_inc is not None:
        amt = float(amount_inc)
        for r in existing_receipts:
            rm = normalize_merchant(r.get("merchant") or r.get("ocr_merchant") or "")
            rd = (r.get("purchased_on") or r.get("ocr_date") or "")[:10]
            ra = r.get("amount_inc")
            if rm == merchant_norm and rd == purchased_on and ra is not None and abs(float(ra) - amt) <= 0.02:
                return STATUS_POSSIBLE_DUP, "Same merchant/date/amount in Field Expenses", r["id"]

        for it in existing_scan_items:
            im = normalize_merchant(it.get("merchant", ""))
            id_ = (it.get("purchased_on") or "")[:10]
            ia = it.get("amount_inc")
            if im == merchant_norm and id_ == purchased_on and ia is not None and abs(float(ia) - amt) <= 0.02:
                return STATUS_POSSIBLE_DUP, "Same merchant/date/amount in this batch", it["id"]

    return STATUS_NEW, "", ""


# ── AI categorisation ─────────────────────────────────────────────────────────

def ai_categorise(
    merchant: str,
    raw_text: str,
    exp_accounts: list[dict],
    *,
    openai_key: str,
    model: str = "gpt-4o-mini",
) -> tuple[str, str]:
    """Ask OpenAI to pick the best Xero expense account. Returns (code, name)."""
    if not openai_key or not exp_accounts:
        return "", ""
    try:
        import openai as _openai
        acct_lines = "\n".join(
            f"  {a.get('Code', '')} — {a.get('Name', '')}"
            for a in exp_accounts[:60]
        )
        prompt = (
            "You are a bookkeeper. Pick the single most appropriate Xero expense "
            "account code for this supplier invoice.\n\n"
            "Powwash-specific account map: supplier identity plus actual line "
            "items/services beats generic account names. Checkatrade, Bark, "
            "Rated People, Google Ads, Meta/Facebook ads, Yell and directory/"
            "lead suppliers are advertising/marketing, not Cleaning. RAC/"
            "breakdown cover is motor vehicle expenses, not insurance. Tender "
            "POS, Stripe, SumUp, Zettle, Square, Worldpay, Cashflows fees and "
            "card-terminal/payment-processing costs are merchant fees/bank "
            "charges, not IT/software. RingGo/parking/car parks, residents "
            "parking permits, visitor permits and CPZ/controlled-parking-zone "
            "charges are vehicle/travel expenses, not Rates. Cleaning-product suppliers such as ECA Cleaning "
            "are materials/job consumables for our exterior cleaning work, not "
            "the Cleaning account. Microsoft, Google Workspace/Cloud, Adobe, "
            "OpenAI, GitHub, Replit and Fly.io are IT/software unless the "
            "document is clearly for card processing. If a supermarket/forecourt "
            "invoice shows fuel: diesel/DERV is van fuel; unleaded/petrol/E10/"
            "E5 is machinery fuel.\n\n"
            f"Merchant / sender: {merchant}\n\n"
            f"OCR text (first 600 chars):\n{raw_text[:600]}\n\n"
            f"Available accounts:\n{acct_lines}\n\n"
            "Reply with ONLY the account code (e.g. 429). No explanation."
        )
        client = _openai.OpenAI(api_key=openai_key)
        resp   = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=20,
            temperature=0,
        )
        code = (resp.choices[0].message.content or "").strip().split()[0].upper()
        for a in exp_accounts:
            if str(a.get("Code", "")).strip().upper() == code:
                return str(a["Code"]).strip(), str(a.get("Name", "")).strip()
    except Exception as exc:
        _log.warning("ai_categorise failed: %s", exc)
    return "", ""


# ── AI invoice gating ─────────────────────────────────────────────────────────
#
# Two cheap, text-only gates keep junk out of the importer and out of the OCR
# pipeline (which is the costly part):
#   1. ai_is_invoice_email — judges the EMAIL (subject + sender + snippet) before
#      any attachment is downloaded or OCR'd, dropping ONLY clearly non-financial
#      mail (marketing / newsletters / social / login / spam). Both gates are
#      RECALL-BIASED ("never miss a bill"): statements, membership/subscription
#      charges and docs addressed to staff/directors by personal name are KEPT.
#   2. ai_validate_invoice_doc — judges the OCR'd DOCUMENT, keeping invoices,
#      bills, statements of account and membership charges; it rejects only clear
#      quotes/estimates/proformas, logos/non-document images, or docs plainly
#      addressed to a DIFFERENT unrelated company.
# Both fail OPEN (keep) when no OpenAI key is configured, so the importer still
# works on the existing filename/MIME heuristics without AI.

def ai_is_invoice_email(
    subject: str,
    sender: str,
    snippet: str,
    *,
    own_names: list[str] | None = None,
    openai_key: str,
    model: str = "gpt-4o-mini",
) -> tuple[bool, str]:
    """Decide whether an email could be *carrying a payable document for us*.

    Recall-biased: keeps invoices, bills, statements and membership/subscription
    charges (incl. those addressed to staff/directors by personal name); drops
    only clearly non-financial mail (marketing / newsletters / social / login /
    spam). Returns (keep, reason). Cheap text-only call run BEFORE downloading
    attachments. Fails open (keep=True) if AI is unavailable or errors.
    """
    if not openai_key:
        return True, ""
    try:
        import openai as _openai
        us = ", ".join(own_names or []) or "our company (Power Wash / Powwash)"
        prompt = (
            "You triage an accounts-payable inbox. Your priority is to NEVER MISS "
            "a supplier invoice, bill or statement — be INCLUSIVE and keep "
            "anything financial.\n\n"
            f"We are: {us}. Documents are often addressed to our staff or "
            "directors by their personal name too — that still counts as us.\n"
            f"From: {sender}\n"
            f"Subject: {subject}\n"
            f"Preview: {snippet[:400]}\n\n"
            "Answer YES if this email could plausibly carry an invoice, bill, "
            "statement of account, membership/subscription charge, or any other "
            "payable document — even if it is worded as 'statement', 'membership "
            "statement' or 'receipt'. When in any doubt, answer YES.\n"
            "ALWAYS answer YES for these (they are financial even if not the word "
            "'invoice'): insurance premiums/policy fees/'outstanding policy fee', "
            "direct debit mandates/instructions/notices, renewal notices, "
            "utility/telecoms/fuel bills, rent/service charges, and any "
            "email whose subject mentions an amount owed or 'outstanding'.\n"
            "Answer NO ONLY when it is clearly NOT a financial document at all: "
            "marketing/newsletters, promotions, social-media or login/security "
            "notifications, delivery/dispatch tracking with no invoice, or pure "
            "spam. When in any doubt, answer YES — it is far better to surface a "
            "borderline document (the admin can ignore it) than to miss a real "
            "bill.\n"
            "Reply with exactly: YES or NO, then a 3-6 word reason. "
            "Example: 'NO — marketing newsletter'."
        )
        client = _openai.OpenAI(api_key=openai_key)
        resp   = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=20,
            temperature=0,
        )
        out = (resp.choices[0].message.content or "").strip()
        keep = out.upper().lstrip().startswith("Y")
        reason = out.split("—", 1)[-1].strip() if "—" in out else out
        return keep, reason[:120]
    except Exception as exc:
        _log.warning("ai_is_invoice_email failed: %s", exc)
        return True, ""


def ai_validate_invoice_doc(
    merchant: str,
    raw_text: str,
    *,
    own_names: list[str] | None = None,
    openai_key: str,
    model: str = "gpt-4o-mini",
) -> tuple[bool, str]:
    """Confirm an OCR'd document is a payable financial doc we should record.

    Recall-biased: keeps invoices, bills, statements of account and membership
    charges (incl. those addressed to staff/directors by personal name); rejects
    only clear quotes/estimates/proformas, logos/non-document images, or docs
    plainly addressed to a DIFFERENT unrelated company. Returns (is_invoice,
    reason). Cheap text-only call run AFTER OCR. Fails open (is_invoice=True) if
    AI is unavailable or errors.
    """
    if not openai_key:
        return True, ""
    text = (raw_text or "").strip()
    if len(text) < 12:
        # Almost no readable text. Small logos/icons are already dropped by the
        # attachment size/inline filter before OCR, so anything that reaches here
        # with no text is more likely a real document OCR struggled with. Keep it
        # for human review rather than risk dropping a genuine invoice.
        return True, ""
    try:
        import openai as _openai
        us = ", ".join(own_names or []) or "Power Wash / Powwash"
        prompt = (
            "You are an accounts-payable clerk. Decide if this OCR'd document is a "
            "financial document we should record: an INVOICE, BILL, STATEMENT of "
            "account, or membership/subscription charge. Your priority is to NEVER "
            "MISS a real bill — be INCLUSIVE.\n\n"
            f"We are the buyer: {us}. The document may be addressed to the company "
            "OR to one of our staff/directors by their personal name — treat that "
            "as US.\n"
            f"Detected merchant/sender: {merchant}\n"
            f"OCR text (first 1200 chars):\n{text[:1200]}\n\n"
            "Answer YES for invoices, bills, statements of account, membership/"
            "subscription charges, insurance premiums/policy fees, and direct "
            "debit mandates/notices showing an amount. When in any doubt, YES.\n"
            "Important exception: if an accountant sends a tax computation, tax "
            "return, CT600, SA302, or document telling us how much tax is owed "
            "to HMRC, answer NO unless the document is clearly the accountant's "
            "own invoice/fee note for their professional services.\n"
            "Answer NO ONLY when it is clearly one of: a quote/estimate/proforma "
            "that is not yet owed, a tax computation/tax return rather than a "
            "supplier charge, a pure marketing/advert page, a logo or non-document "
            "image, or a document plainly addressed to a DIFFERENT, unrelated "
            "company (neither us nor our staff). When in any doubt, answer YES.\n"
            "Reply with exactly: YES or NO, then a 3-6 word reason. "
            "Example: 'NO — quotation, not an invoice'."
        )
        client = _openai.OpenAI(api_key=openai_key)
        resp   = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=24,
            temperature=0,
        )
        out = (resp.choices[0].message.content or "").strip()
        is_inv = out.upper().lstrip().startswith("Y")
        reason = out.split("—", 1)[-1].strip() if "—" in out else out
        return is_inv, reason[:120]
    except Exception as exc:
        _log.warning("ai_validate_invoice_doc failed: %s", exc)
        return True, ""


# ── Error translation ─────────────────────────────────────────────────────────

def friendly_gmail_error(exc: Exception) -> str:
    """Turn a raw Gmail/Google API exception into a plain, actionable message."""
    msg = str(exc or "").strip()
    low = msg.lower()
    if (
        "accessnotconfigured" in low
        or "has not been used in project" in low
        or ("gmail api" in low and ("disabled" in low or "not been used" in low))
    ):
        return (
            "Gmail API is not enabled for your Google Cloud project. "
            "Open Google Cloud Console → APIs & Services → Library, search for "
            "\"Gmail API\", click Enable, wait about a minute, then run the scan "
            "again. (Connecting Google grants permission, but the Gmail API must "
            "also be switched on in the project.)"
        )
    if (
        "insufficientpermissions" in low
        or "insufficient authentication scopes" in low
        or "request had insufficient" in low
        or ("scope" in low and "gmail" in low)
    ):
        return (
            "Gmail read-only access has not been approved. Go to Settings → Google, "
            "click Reconnect, approve the Gmail permission, then run the scan again."
        )
    if (
        "invalid_grant" in low
        or "unauthorized" in low
        or "invalid credentials" in low
        or "token has been expired" in low
        or "401" in low
    ):
        return (
            "Your Google sign-in has expired. Go to Settings → Google, click "
            "Reconnect, then run the scan again."
        )
    if "ratelimit" in low or "rate limit" in low or "quota" in low or "userratelimit" in low:
        return "Gmail rate limit/quota reached. Wait a few minutes and run the scan again."
    return (msg[:300] or "Unknown error during scan.")


# ── Main scan entry point ─────────────────────────────────────────────────────

def scan_email_batch(
    batch_id: str,
    db_path: str,
    *,
    gmail_creds,
    own_names: list[str],
    own_domains: list[str],
    exp_accounts: list[dict],
    vat_rate: float = 20.0,
    openai_key: str = "",
    receipt_service,
    max_messages: int = 500,
    progress_cb=None,
) -> dict[str, Any]:
    """Scan Gmail for invoice attachments and populate email_scan_items.

    Never modifies Gmail messages — uses gmail.readonly scope only.

    Returns a summary dict: found / own_company / duplicate / no_invoice / new / error.
    """
    from ..gmail_client import GmailClient, is_invoice_attachment as _inv_att

    batch = get_batch(db_path, batch_id)
    if not batch:
        return {"error": "Batch not found"}

    date_from = batch.get("date_from", "")
    date_to   = batch.get("date_to",   "")

    def _cb(msg: str) -> None:
        if progress_cb:
            try:
                progress_cb(msg)
            except Exception:
                pass

    try:
        _cb("Connecting to Gmail…")
        client   = GmailClient(gmail_creds)
        messages = client.search_messages(date_from, date_to, max_results=max_messages)
        total    = len(messages)
        _cb(f"Found {total} email(s) with attachments. Analysing…")

        existing_receipts: list[dict] = list_all_receipts(db_path)
        existing_items:    list[dict] = []   # grows per item so in-batch dedup works

        counts: dict[str, int] = {
            "found": 0, "own_company": 0, "duplicate": 0,
            "no_invoice": 0, "new": 0, "error": 0, "no_account": 0,
            "not_invoice": 0, "skipped_email": 0,
        }
        seq = 0

        def _record_skip(status: str, reason: str, *, msg_id: str,
                         thread_id: str, from_addr: str, from_name: str,
                         subject: str, email_date: str) -> None:
            """Store a lightweight, reviewable placeholder for an email we are
            NOT importing (AI-skipped / attachments filtered). Own-company
            outgoing invoices are NOT recorded — they are dropped entirely.
            Nothing is downloaded or OCR'd; the admin can force a full scan of
            the message later via 'Scan anyway'. Placeholders have no amount so
            they never take part in card-feed matching or imports."""
            nonlocal seq
            try:
                it = create_item(
                    db_path,
                    batch_id=batch_id,
                    seq=seq,
                    message_id=msg_id,
                    thread_id=thread_id,
                    sender_from=from_addr,
                    sender_name=from_name,
                    subject=subject,
                    email_date=email_date,
                    status=status,
                    merchant=from_name or from_addr,
                    purchased_on=(email_date or "")[:10],
                    dup_reason=reason,
                )
                existing_items.append(it)
                seq += 1
            except Exception as exc:
                _log.warning("record skip %s: %s", msg_id, exc)

        for idx, stub in enumerate(messages):
            msg_id = stub["id"]

            try:
                meta = client.get_message_metadata(msg_id)
            except Exception as exc:
                _log.warning("Gmail meta %s: %s", msg_id, exc)
                counts["error"] += 1
                continue

            from_addr  = meta.get("from", "")
            from_name  = meta.get("from_name", "")
            subject    = meta.get("subject", "")
            email_date = meta.get("date_iso", "")
            thread_id  = meta.get("thread_id", "")

            if is_own_company_sender(
                from_addr, from_name, subject,
                own_names=own_names, own_domains=own_domains,
            ):
                # Our OWN outgoing invoices (sent by us, sent on our behalf, or
                # customer reply threads like "Re: Invoice INV-5691 from Pow
                # Services Limited ...") are dead space — dropped entirely, not
                # stored as reviewable placeholders. Only the summary counts
                # them. Supplier invoices say "for/to [us]", never "from [us]",
                # so they are unaffected.
                counts["own_company"] += 1
                continue

            # Gate 1 — cheap, text-only AI triage of the EMAIL before we download
            # or OCR anything. Recall-biased: drops ONLY clearly non-financial
            # mail (marketing / newsletters / social / login / spam); keeps
            # statements & membership charges so real bills are never missed.
            keep_email, _email_reason = ai_is_invoice_email(
                subject, from_addr or from_name, meta.get("snippet", ""),
                own_names=own_names, openai_key=openai_key,
            )
            if not keep_email:
                counts["skipped_email"] += 1
                _record_skip(
                    STATUS_SKIPPED_EMAIL,
                    "AI skipped this email: "
                    + (_email_reason or "judged not financial"),
                    msg_id=msg_id, thread_id=thread_id, from_addr=from_addr,
                    from_name=from_name, subject=subject, email_date=email_date,
                )
                continue

            try:
                attachments = client.get_attachments(msg_id)
            except Exception as exc:
                _log.warning("Gmail attachments %s: %s", msg_id, exc)
                counts["error"] += 1
                continue

            msg_any_item = False   # created a real (OCR'd) item for this email
            msg_any_dup  = False   # skipped only because already imported

            for att in attachments:
                fname = att.get("filename", "unknown")
                mime  = att.get("mime_type", "")
                # Gmail sometimes reports application/octet-stream for PDFs/images
                # when the sender's mail client omits a Content-Type. Infer the real
                # type from the filename extension so Document AI can OCR it.
                if (mime or "").lower().strip() == "application/octet-stream" and "." in fname:
                    _ext = ("." + fname.rsplit(".", 1)[-1]).lower()
                    mime = {
                        ".pdf":  "application/pdf",
                        ".jpg":  "image/jpeg", ".jpeg": "image/jpeg",
                        ".png":  "image/png",
                        ".tif":  "image/tiff", ".tiff": "image/tiff",
                        ".gif":  "image/gif",
                        ".bmp":  "image/bmp",
                        ".webp": "image/webp",
                        ".docx": _DOCX_MIME,
                    }.get(_ext, mime)
                data  = att.get("data_bytes", b"")

                if not _inv_att(
                    fname, mime,
                    inline=bool(att.get("inline")),
                    size=att.get("size") if att.get("size") is not None else len(data),
                ) or not data:
                    counts["no_invoice"] += 1
                    continue

                if message_already_scanned(db_path, msg_id, fname):
                    counts["duplicate"] += 1
                    msg_any_dup = True
                    continue

                sha = sha256_hex(data)

                # OCR — Word .docx invoices can't go through Document AI, so we
                # extract their text directly; everything else uses Document AI.
                if is_word_doc(fname, mime):
                    try:
                        if len(data) > 12 * 1024 * 1024:
                            raise ValueError("file too large")
                        _wtext = extract_docx_text(data)
                        _store_err = ""
                        try:
                            _stored = receipt_service.store_file(
                                file_bytes=data, filename=fname
                            )
                        except Exception as sexc:
                            _stored = ""
                            _store_err = f"Could not store Word document: {sexc}"[:200]
                            _log.warning("DOCX store %s/%s: %s", msg_id, fname, sexc)
                        result = {
                            "merchant": "", "total": None, "net": None, "tax": None,
                            "date": "", "stored_file": _stored, "raw_text": _wtext,
                            "ocr_error": (_store_err or ("" if _wtext.strip()
                                          else "Could not read Word document")),
                        }
                    except Exception as exc:
                        _log.warning("DOCX %s/%s: %s", msg_id, fname, exc)
                        result = {
                            "merchant": "", "total": None, "net": None, "tax": None,
                            "date": "", "stored_file": "", "raw_text": "",
                            "ocr_error": str(exc)[:200],
                        }
                else:
                    try:
                        result = receipt_service.analyze_upload(
                            file_bytes=data, filename=fname, mime_type=mime
                        )
                    except Exception as exc:
                        _log.warning("OCR %s/%s: %s", msg_id, fname, exc)
                        result = {
                            "merchant": "", "total": None, "net": None, "tax": None,
                            "date": "", "stored_file": "", "raw_text": "",
                            "ocr_error": str(exc)[:200],
                        }

                merchant    = (result.get("merchant") or "").strip()
                merchant    = derive_supplier_merchant(
                    merchant, from_name=from_name, from_addr=from_addr,
                    subject=subject, own_names=own_names,
                )
                raw_text    = result.get("raw_text") or result.get("ocr_raw") or ""
                stored_file = result.get("stored_file", "")
                ocr_error   = result.get("ocr_error", "")

                inc, ex, vat = reconcile_email_amounts_from_text(
                    result.get("total") or result.get("amount_inc"),
                    result.get("net")   or result.get("amount_ex"),
                    result.get("tax")   or result.get("vat_amount"),
                    raw_text,
                    vat_rate=vat_rate,
                )
                # Fallback: if Document AI couldn't parse a structured amount
                # (e.g. non-standard council/parking receipts), try to find the
                # largest currency-like number in the raw OCR text.
                if inc is None and raw_text and not ocr_error:
                    _amounts = []
                    for _m in re.findall(r'\b(\d{1,6}(?:,\d{3})*\.\d{1,2}|\d{1,5}\.\d{1,2})\b', raw_text):
                        try:
                            _v = float(_m.replace(",", ""))
                            if _v >= 0.50:
                                _amounts.append(_v)
                        except ValueError:
                            pass
                    if _amounts:
                        inc = max(_amounts)
                purchased_on  = (
                    result.get("date") or result.get("purchased_on") or email_date or ""
                )[:10]
                merchant_norm = normalize_merchant(merchant or from_name)

                # Dedup
                status, dup_reason, match_id = dedup_against_receipts(
                    sha, merchant_norm, purchased_on, inc,
                    existing_receipts, existing_items,
                )

                # Gate 2 — for genuinely new items, confirm the OCR'd DOCUMENT is
                # a payable financial doc (invoice / bill / statement / membership
                # charge), rejecting only clear quotes/estimates, logos, or docs
                # addressed to a different unrelated company. Skipped for
                # duplicates (already known) to save AI cost.
                final_status = status
                cat_code, cat_name = "", ""
                if status == STATUS_NEW:
                    why_not = tax_computation_not_supplier_invoice(
                        merchant or from_name, raw_text
                    )
                    is_invoice = not bool(why_not)
                    if is_invoice:
                        is_invoice, why_not = ai_validate_invoice_doc(
                            merchant or from_name, raw_text,
                            own_names=own_names, openai_key=openai_key,
                        )
                    if not is_invoice:
                        final_status = STATUS_NOT_INVOICE
                        dup_reason = why_not or "AI: not an invoice to Power Wash"
                        counts["not_invoice"] += 1

                # Categorise (history first, then AI) only for confirmed invoices
                if final_status == STATUS_NEW:
                    cat_code, cat_name = learned_categorise(db_path, merchant or from_name)
                    if not cat_code:
                        cat_code, cat_name = rule_based_categorise(
                            merchant or from_name, raw_text, exp_accounts
                        )
                    if not cat_code:
                        cat_code, cat_name = smart_categorise(
                            merchant_norm,
                            exp_receipts=existing_receipts,
                            email_items=existing_items,
                        )
                    if not cat_code:
                        cat_code, cat_name = ai_categorise(
                            merchant or from_name, raw_text, exp_accounts,
                            openai_key=openai_key,
                        )
                    if not cat_code:
                        final_status = STATUS_NO_ACCOUNT

                item = create_item(
                    db_path,
                    batch_id=batch_id,
                    seq=seq,
                    message_id=msg_id,
                    thread_id=thread_id,
                    sender_from=from_addr,
                    sender_name=from_name,
                    subject=subject,
                    email_date=email_date,
                    attachment_name=fname,
                    attachment_mime=mime,
                    status=final_status,
                    merchant=merchant or from_name,
                    purchased_on=purchased_on,
                    amount_inc=inc,
                    amount_ex=ex,
                    vat_amount=vat,
                    currency="GBP",
                    category_account_code=cat_code,
                    category_account_name=cat_name,
                    dup_reason=dup_reason,
                    match_receipt_id=match_id,
                    stored_file=stored_file,
                    ocr_raw=raw_text[:2000],
                    ocr_error=ocr_error,
                )
                existing_items.append(item)
                seq += 1
                msg_any_item = True
                counts["found"] += 1
                if final_status == STATUS_NEW:
                    counts["new"] += 1
                elif final_status == STATUS_NO_ACCOUNT:
                    counts["no_account"] += 1
                elif final_status in (STATUS_DUPLICATE, STATUS_POSSIBLE_DUP):
                    counts["duplicate"] += 1

            # The email survived both the own-company and AI gates but NOTHING
            # was scanned from it (every attachment looked like a photo / logo /
            # icon, or had no data). Record a reviewable placeholder so a real
            # invoice with an odd attachment is never silently lost.
            if not msg_any_item and not msg_any_dup:
                counts["skipped_email"] += 1
                _record_skip(
                    STATUS_SKIPPED_EMAIL,
                    "Attachments looked like photos/logos/icons — nothing was "
                    "scanned. Use 'Scan anyway' if this is a real invoice.",
                    msg_id=msg_id, thread_id=thread_id, from_addr=from_addr,
                    from_name=from_name, subject=subject, email_date=email_date,
                )

            if (idx + 1) % 5 == 0:
                _cb(f"Processed {idx + 1}/{total} emails…")
                # Persist partial progress so the polling results page shows the
                # scan moving instead of appearing frozen on a long mailbox.
                update_batch(
                    db_path, batch_id,
                    status="processing",
                    total_found=counts["found"],
                    summary_json=json.dumps({**counts, "scanned": idx + 1, "total": total}),
                )

        update_batch(
            db_path, batch_id,
            status="ready",
            total_found=counts["found"],
            summary_json=json.dumps(counts),
        )
        _cb(f"Done — {counts['new']} new, {counts['duplicate']} duplicate(s) found.")
        return counts

    except Exception as exc:
        _log.exception("scan_email_batch %s: %s", batch_id, exc)
        friendly = friendly_gmail_error(exc)
        update_batch(
            db_path, batch_id,
            status="error",
            summary_json=json.dumps({"error": friendly}),
        )
        return {"error": friendly}


# ── Force-rescan a single skipped email ──────────────────────────────────────

def rescan_message(
    batch_id: str,
    item_id: str,
    db_path: str,
    *,
    gmail_creds,
    exp_accounts: list[dict],
    vat_rate: float = 20.0,
    openai_key: str = "",
    receipt_service,
    own_names: list[str] | None = None,
) -> tuple[int, str]:
    """Force-process ONE email's attachments — used by 'Scan anyway'.

    Bypasses the AI email gate, the AI document gate AND the photo/logo
    attachment filter (a human has explicitly said "this is an invoice"), but
    still runs OCR, duplicate detection and account categorisation. The first
    processed attachment replaces the placeholder item in place; extra
    attachments become new items. Returns (n_processed, error_message).
    """
    from ..gmail_client import GmailClient

    items  = list_items(db_path, batch_id)
    target = next((i for i in items if i["id"] == item_id), None)
    if not target:
        return 0, "Item not found in this scan."
    msg_id = (target.get("message_id") or "").strip()
    if not msg_id:
        return 0, "This item has no Gmail message linked to it."

    try:
        client      = GmailClient(gmail_creds)
        meta        = client.get_message_metadata(msg_id)
        attachments = client.get_attachments(msg_id)
    except Exception as exc:
        _log.warning("rescan %s: %s", msg_id, exc)
        return 0, friendly_gmail_error(exc)

    from_addr  = meta.get("from", "")      or target.get("sender_from", "")
    from_name  = meta.get("from_name", "") or target.get("sender_name", "")
    email_date = meta.get("date_iso", "")  or target.get("email_date", "")

    # Relaxed filter: any PDF / image / Word attachment with data counts —
    # photo-style filenames, inline parts and small files are all allowed.
    usable: list[tuple[str, str, bytes]] = []
    for att in attachments:
        fname = att.get("filename", "unknown")
        mime  = (att.get("mime_type") or "").lower().split(";")[0].strip()
        data  = att.get("data_bytes", b"")
        if not data:
            continue
        ext = ("." + fname.rsplit(".", 1)[-1]).lower() if "." in fname else ""
        if mime == "application/octet-stream" and ext:
            mime = {
                ".pdf":  "application/pdf",
                ".jpg":  "image/jpeg", ".jpeg": "image/jpeg",
                ".png":  "image/png",
                ".tif":  "image/tiff", ".tiff": "image/tiff",
                ".gif":  "image/gif",
                ".webp": "image/webp",
                ".docx": _DOCX_MIME,
            }.get(ext, mime)
        if (ext in _INVOICE_EXTS or mime in _INVOICE_MIMES
                or is_word_doc(fname, mime)):
            usable.append((fname, mime, data))

    if not usable:
        update_item(
            db_path, item_id,
            dup_reason="Rescanned — this email has no readable PDF/image/Word "
                       "attachment, so there is nothing to import from it.",
        )
        return 0, ""

    existing_receipts = list_all_receipts(db_path)
    existing_items    = [i for i in items if i["id"] != item_id]
    max_seq = max([int(i.get("seq") or 0) for i in items] or [0])

    n_done  = 0
    status_counts: dict[str, int] = {}
    replaced_placeholder = False
    for fname, mime, data in usable:
        sha = sha256_hex(data)

        if is_word_doc(fname, mime):
            try:
                if len(data) > 12 * 1024 * 1024:
                    raise ValueError("file too large")
                _wtext = extract_docx_text(data)
                try:
                    _stored = receipt_service.store_file(
                        file_bytes=data, filename=fname)
                    _store_err = ""
                except Exception as sexc:
                    _stored, _store_err = "", f"Could not store Word document: {sexc}"[:200]
                result = {
                    "merchant": "", "total": None, "net": None, "tax": None,
                    "date": "", "stored_file": _stored, "raw_text": _wtext,
                    "ocr_error": (_store_err or ("" if _wtext.strip()
                                  else "Could not read Word document")),
                }
            except Exception as exc:
                result = {"merchant": "", "total": None, "net": None,
                          "tax": None, "date": "", "stored_file": "",
                          "raw_text": "", "ocr_error": str(exc)[:200]}
        else:
            try:
                result = receipt_service.analyze_upload(
                    file_bytes=data, filename=fname, mime_type=mime)
            except Exception as exc:
                result = {"merchant": "", "total": None, "net": None,
                          "tax": None, "date": "", "stored_file": "",
                          "raw_text": "", "ocr_error": str(exc)[:200]}

        merchant    = (result.get("merchant") or "").strip()
        merchant    = derive_supplier_merchant(
            merchant, from_name=from_name, from_addr=from_addr,
            subject=target.get("subject", ""), own_names=own_names,
        )
        raw_text    = result.get("raw_text") or result.get("ocr_raw") or ""
        stored_file = result.get("stored_file", "")
        ocr_error   = result.get("ocr_error", "")

        inc, ex, vat = reconcile_email_amounts_from_text(
            result.get("total") or result.get("amount_inc"),
            result.get("net")   or result.get("amount_ex"),
            result.get("tax")   or result.get("vat_amount"),
            raw_text,
            vat_rate=vat_rate,
        )
        if inc is None and raw_text and not ocr_error:
            _amounts = []
            for _m in re.findall(r'\b(\d{1,6}(?:,\d{3})*\.\d{1,2}|\d{1,5}\.\d{1,2})\b', raw_text):
                try:
                    _v = float(_m.replace(",", ""))
                    if _v >= 0.50:
                        _amounts.append(_v)
                except ValueError:
                    pass
            if _amounts:
                inc = max(_amounts)

        purchased_on = (
            result.get("date") or result.get("purchased_on") or email_date or ""
        )[:10]
        merchant_norm = normalize_merchant(merchant or from_name)

        status, dup_reason, match_id = dedup_against_receipts(
            sha, merchant_norm, purchased_on, inc,
            existing_receipts, existing_items,
        )

        final_status = status
        cat_code, cat_name = "", ""
        if status == STATUS_NEW:
            cat_code, cat_name = learned_categorise(db_path, merchant or from_name)
            if not cat_code:
                cat_code, cat_name = rule_based_categorise(
                    merchant or from_name, raw_text, exp_accounts
                )
            if not cat_code:
                cat_code, cat_name = smart_categorise(
                    merchant_norm,
                    exp_receipts=existing_receipts,
                    email_items=existing_items,
                )
            if not cat_code:
                cat_code, cat_name = ai_categorise(
                    merchant or from_name, raw_text, exp_accounts,
                    openai_key=openai_key,
                )
            if not cat_code:
                final_status = STATUS_NO_ACCOUNT

        fields = dict(
            attachment_name=fname,
            attachment_mime=mime,
            status=final_status,
            merchant=merchant or from_name,
            purchased_on=purchased_on,
            amount_inc=inc,
            amount_ex=ex,
            vat_amount=vat,
            category_account_code=cat_code,
            category_account_name=cat_name,
            dup_reason=dup_reason or "Force-scanned by admin",
            match_receipt_id=match_id,
            stored_file=stored_file,
            ocr_raw=raw_text[:2000],
            ocr_error=ocr_error,
        )
        if not replaced_placeholder:
            it = update_item(db_path, item_id, **fields)
            replaced_placeholder = True
        else:
            max_seq += 1
            it = create_item(
                db_path, batch_id=batch_id, seq=max_seq,
                message_id=msg_id,
                thread_id=target.get("thread_id", ""),
                sender_from=from_addr,
                sender_name=from_name,
                subject=target.get("subject", ""),
                email_date=email_date,
                currency="GBP",
                **fields,
            )
        if it:
            existing_items.append(it)
        n_done += 1
        status_counts[final_status] = status_counts.get(final_status, 0) + 1

    # Keep the batch summary roughly honest.
    try:
        b = get_batch(db_path, batch_id) or {}
        summ = b.get("summary") or {}
        if isinstance(summ, str):
            summ = json.loads(summ)
        summ["found"] = int(summ.get("found") or 0) + n_done
        for _st, _n in status_counts.items():
            summ[_st] = int(summ.get(_st) or 0) + _n
        if target.get("status") == STATUS_SKIPPED_EMAIL:
            summ["skipped_email"] = max(0, int(summ.get("skipped_email") or 0) - 1)
        elif target.get("status") == STATUS_OWN_COMPANY:
            summ["own_company"] = max(0, int(summ.get("own_company") or 0) - 1)
        update_batch(
            db_path, batch_id,
            total_found=int(b.get("total_found") or 0) + n_done,
            summary_json=json.dumps(summ),
        )
    except Exception as exc:
        _log.warning("rescan summary update %s: %s", batch_id, exc)

    return n_done, ""


# ── Import (write to expense_receipts) ───────────────────────────────────────

def import_batch_items(
    batch_id: str,
    db_path: str,
    *,
    default_engineer_id: int,
) -> int:
    """Write STATUS_NEW email_scan_items into expense_receipts.

    Returns number of rows imported.
    """
    items    = list_items(db_path, batch_id)
    imported = 0
    existing_receipts = list_all_receipts(db_path, limit=1_000_000)
    for it in items:
        if it["status"] != STATUS_NEW:
            continue
        try:
            fresh_status, fresh_reason, fresh_match_id = dedup_against_receipts(
                "",
                normalize_merchant(it.get("merchant", "")),
                (it.get("purchased_on") or "")[:10],
                it.get("amount_inc"),
                existing_receipts,
                [],
            )
            if fresh_status != STATUS_NEW:
                update_item(
                    db_path,
                    it["id"],
                    status=fresh_status,
                    dup_reason=(
                        fresh_reason
                        or "Duplicate found during final import check"
                    ),
                    match_receipt_id=fresh_match_id,
                )
                continue

            rec = create_receipt(
                db_path,
                engineer_id=default_engineer_id,
                merchant=it.get("merchant", ""),
                purchased_on=it.get("purchased_on", ""),
                amount_inc=it.get("amount_inc"),
                amount_ex=it.get("amount_ex"),
                vat_amount=it.get("vat_amount"),
                currency=it.get("currency", "GBP"),
                ocr_merchant=it.get("merchant", ""),
                ocr_amount=it.get("amount_inc"),
                ocr_date=it.get("purchased_on", ""),
                ocr_raw=it.get("ocr_raw", ""),
                ocr_error=it.get("ocr_error", ""),
                stored_file=it.get("stored_file", ""),
                filename=it.get("attachment_name", ""),
                mime_type=it.get("attachment_mime", ""),
                category_account_code=it.get("category_account_code", ""),
                category_account_name=it.get("category_account_name", ""),
                status="pending_review",
            )
            if rec:
                existing_receipts.insert(0, rec)
            update_item(db_path, it["id"], status=STATUS_IMPORTED)
            imported += 1
        except Exception as exc:
            _log.warning("import item %s: %s", it["id"], exc)

    update_batch(db_path, batch_id, status="done")
    return imported
